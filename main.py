from tqdm import tqdm
from pdf2image import convert_from_path
import pytesseract
from PIL import Image

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from pathlib import Path
import tempfile
import sys


def pdf_to_word(pdf_path:str, output_dir:str, lang='fas', **kwargs):
    """ 
    convert pdf to word
    """
    pdf_path = Path(pdf_path)
    pdf_name = pdf_path.stem
    pages = convert_from_path(pdf_path, **kwargs)
    texts = []

    print(f'PDF is preparing to convert into document [#{len(pages)} pages]')
    for i, page in tqdm(enumerate(pages), position=0):

        with tempfile.TemporaryDirectory() as img_dir:
            img_name = f'{pdf_name}-{i+1}.jpg'
            img_path = Path(img_dir) / img_name

            page.save(img_path, 'JPEG')
            text = pytesseract.image_to_string(Image.open(img_path), lang=lang)
            texts.append(text)

    document = Document()
    style_normal = document.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.rtl = True

    style_h1 = document.styles['Heading 1']
    font = style_h1.font
    font.name = 'Arial'
    font.rtl = True

    for i, text in tqdm(enumerate(texts), position=0):
        heading = document.add_heading(f'صفحه: {i+1}', level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        heading.style = document.styles['Heading 1']

        paragraph = document.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraph.style = document.styles['Normal']

    output_path = Path(output_dir) / f'{pdf_name}.docx'
    document.save(output_path)
    print(f'Done! Your document can be find here "{output_path}"')


if __name__ == '__main__':
    args = sys.argv[1:]
    if len(args) >= 1:
        pdf_to_word(args[0], args[1])
    else:
        pdf_path = input('PDF path for convertion: ')
        output_dir = input('Output directory: ')
        pdf_to_word(pdf_path, output_dir)
